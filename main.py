from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Form
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import json
import requests
import os
import sys
import uuid
import logging
from datetime import datetime
import time
import aiofiles
from pathlib import Path
from dotenv import load_dotenv

PROJECT_ROOT = Path(__file__).parent
ENV_FILE = PROJECT_ROOT / "config" / ".env"

if ENV_FILE.exists():
    load_dotenv(ENV_FILE)
else:
    logger_setup = logging.getLogger(__name__)
    logger_setup.warning(f"Archivo .env no encontrado en: {ENV_FILE}")

# Añadir la raíz del proyecto al path para que el paquete `src` pueda importarse como top-level
sys.path.insert(0, str(PROJECT_ROOT))

from src.transcriber import AudioTranscriber
from src.diarizer import SpeakerDiarizer
from src.video_converter import VideoConverter, is_video_file
from src.utils import (
    align_transcription_with_diarization,
    format_transcript,
    get_speaker_statistics,
    renumber_speakers,
    ensure_upload_dirs,
)
# Worker in-process will be importado en startup para evitar fallos de importación
worker_module = None

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Audio Transcription & Diarization API",
    description="API para transcribir audio e identificar hablantes",
    version="1.0.0"
)


def _worker_file_head():
    """Comprueba si /app/src/worker.py existe y devuelve las primeras líneas o None.

    Esto se separa para evitar problemas de scope con Path dentro de la función de startup.
    """
    try:
        p_check = Path(__file__).resolve().parent / 'src' / 'worker.py'
        if p_check.exists():
            try:
                with open(p_check, 'r', encoding='utf-8') as fh:
                    first_lines = ''.join([next(fh) for _ in range(5)])
                return (p_check, first_lines)
            except Exception:
                return (p_check, None)
        else:
            return (p_check, None)
    except Exception:
        return (None, None)


@app.on_event('startup')
def start_worker_thread():
    """Lanza el worker en un hilo daemon para procesar jobs en `uploads/jobs/`.

    Esta es una solución práctica cuando no se despliega un worker separado.
    Requiere que la instancia tenga recursos suficientes (tú ya subiste a 16GB).
    """
    import threading
    global worker_module

    # Permitir desactivar el worker embebido mediante variable de entorno (útil durante pruebas locales)
    if os.getenv('DISABLE_EMBEDDED_WORKER', '').lower() in ('1', 'true', 'yes'):
        logger.info('DISABLE_EMBEDDED_WORKER set -> no se iniciará el worker embebido')
        return

    # Importar el módulo del worker aquí y registrar cualquier excepción de import
    if worker_module is None:
        # Diagnostic: verificar si el archivo src/worker.py existe dentro del contenedor
        p_check, first_lines = _worker_file_head()
        if p_check is None:
            logger.info("DEBUG: no se pudo determinar la ruta de worker.py")
        else:
            try:
                if p_check.exists():
                    logger.info(f"DEBUG: worker.py existe en la imagen en: {p_check} (size={p_check.stat().st_size} bytes)")
                    if first_lines:
                        logger.info(f"DEBUG: primeras líneas de worker.py:\n{first_lines}")
                    else:
                        logger.info("DEBUG: no se pudieron leer las primeras líneas de worker.py")
                else:
                    logger.info("DEBUG: worker.py NO existe en la imagen /app/src/")
            except Exception:
                logger.exception("DEBUG: error comprobando existencia de worker.py")
        try:
            import importlib
            worker_module = importlib.import_module('src.worker')
        except Exception as e:
            logger.exception(f"No se pudo importar module worker en startup: {e}")
            # Fallback: intentar cargar por path (útil cuando el paquete no está en sys.path)
            try:
                import importlib.util
                p = Path(__file__).resolve().parent / 'src' / 'worker.py'
                if p.exists():
                    spec = importlib.util.spec_from_file_location('embedded_worker', str(p))
                    wm = importlib.util.module_from_spec(spec)
                    spec.loader.exec_module(wm)
                    worker_module = wm
                    logger.info(f"Worker cargado desde archivo: {p}")
                else:
                    worker_module = None
            except Exception as e2:
                logger.exception(f"Fallback para cargar worker desde archivo falló: {e2}")
                worker_module = None

            # Si no había worker en la imagen, intentar descargarlo desde GitHub (raw) y cargarlo dinámicamente.
            if worker_module is None:
                try:
                    raw_url = "https://raw.githubusercontent.com/Leocix/Transcripci-n-de-Audio/master/src/worker.py"
                    logger.info(f"Intentando descargar worker desde: {raw_url}")
                    # Usar requests con timeout corto
                    resp = requests.get(raw_url, timeout=10)
                    if resp.status_code == 200 and resp.text:
                        # Escribir en un archivo temporal dentro del proyecto (no commit)
                        tmp_path = Path(__file__).resolve().parent / 'src' / '_downloaded_worker.py'
                        try:
                            tmp_path.parent.mkdir(parents=True, exist_ok=True)
                            tmp_path.write_text(resp.text, encoding='utf-8')
                            spec = importlib.util.spec_from_file_location('downloaded_worker', str(tmp_path))
                            wm2 = importlib.util.module_from_spec(spec)
                            spec.loader.exec_module(wm2)
                            worker_module = wm2
                            logger.info(f"Worker dinámico cargado desde {raw_url} y guardado en {tmp_path}")
                        except Exception as e3:
                            logger.exception(f"Error al escribir/cargar worker descargado: {e3}")
                            worker_module = None
                    else:
                        logger.warning(f"No se pudo descargar worker desde raw URL, status_code={resp.status_code}")
                except Exception as e4:
                    logger.exception(f"Excepción al intentar descargar worker desde GitHub: {e4}")

    if worker_module is None:
        logger.info("Worker module no disponible; no se iniciará el worker embebido.")
        return

    def _run():
        try:
            logger.info("Embedded worker thread iniciando main_loop()")
            # Inyectar funciones de actualización en el worker
            if hasattr(worker_module, 'set_job_update_functions'):
                worker_module.set_job_update_functions(_update_job)
            worker_module.main_loop()
        except Exception as e:
            logger.exception(f"Worker embebido terminó con error: {e}")

    t = threading.Thread(target=_run, name='embedded-worker', daemon=True)
    t.start()
    logger.info("Hilo del worker embebido lanzado")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

WEB_DIR = PROJECT_ROOT / "web"
UPLOAD_DIR = os.getenv("UPLOAD_DIR", str(PROJECT_ROOT / "uploads"))
WHISPER_MODEL = os.getenv("WHISPER_MODEL", "base")
HF_TOKEN = os.getenv("HF_TOKEN")
MAX_FILE_SIZE = int(os.getenv("MAX_FILE_SIZE", 100 * 1024 * 1024))

# Asegurar que las carpetas de uploads y subdirectorios existan (jobs, results, jobs/failed)
paths = ensure_upload_dirs(UPLOAD_DIR)
JOBS_DIR = paths.get('jobs_dir')
RESULTS_DIR = paths.get('results_dir')

transcriber = None
diarizer = None
video_converter = None

# In-memory job store: job_id -> status dict
jobs = {}


def _create_job(job_id: str, meta: dict = None):
    jobs[job_id] = {
        "job_id": job_id,
        "state": "pending",
        "progress": 0,
        "message": "pending",
        "start_time": time.time(),
        "last_update": time.time(),
        "meta": meta or {},
        "result": None,
        "error": None
    }


def _update_job(job_id: str, progress: int = None, message: str = None, result: dict = None, state: str = None, error: str = None):
    job = jobs.get(job_id)
    if not job:
        return
    now = time.time()
    if progress is not None:
        job["progress"] = max(0, min(100, int(progress)))
    if message is not None:
        job["message"] = message
    job["last_update"] = now
    if state is not None:
        job["state"] = state
    if result is not None:
        job["result"] = result
    if error is not None:
        job["error"] = error


def _estimate_remaining(job: dict):
    # Estimar tiempo restante usando elapsed / progress
    elapsed = time.time() - job.get("start_time", time.time())
    progress = job.get("progress", 0)
    if progress <= 0:
        return None
    remaining = elapsed * (100 - progress) / progress
    return int(remaining)


def run_transcription_job(job_id: str, file_path: str, language: Optional[str], task: str, download_format: Optional[str]):
    try:
        _update_job(job_id, state="running", message="iniciando", progress=1)
        logger.info(f"[JOB {job_id}] Cargando modelo para transcripción")
        transcriber_instance = get_transcriber()

        _update_job(job_id, progress=5, message="transcribiendo")
        result = transcriber_instance.transcribe(file_path, language=language, task=task)
        _update_job(job_id, progress=60, message="transcripción completa")

        _update_job(job_id, progress=80, message="procesando resultados")
        # Guardar resultados
        job_result = {
            "text": result.get("text"),
            "language": result.get("language"),
            "segments": result.get("segments"),
            "model": result.get("model")
        }

        # Export si aplica
        if download_format in ("docx", "pdf"):
            out_name = f"{job_id}.{download_format}"
            out_path = os.path.join(UPLOAD_DIR, out_name)
            if download_format == "docx":
                _save_docx(job_result["text"], out_path)
            else:
                _save_pdf(job_result["text"], out_path)
            job_result["download_path"] = out_path

        _update_job(job_id, progress=100, message="completado", state="done", result=job_result)
        logger.info(f"[JOB {job_id}] Completado")
    except Exception as e:
        logger.error(f"[JOB {job_id}] Error: {e}")
        _update_job(job_id, state="error", error=str(e), message="error")
    finally:
        # Eliminar archivo temporal si existe
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"[JOB {job_id}] Archivo temporal eliminado: {file_path}")
        except Exception as e:
            logger.warning(f"[JOB {job_id}] No se pudo eliminar archivo temporal: {e}")


@app.get("/status/{job_id}", tags=["Jobs"])
async def job_status(job_id: str):
    # Primero intentar resolver desde el store en memoria
    job = jobs.get(job_id)

    if not job:
        # Intentar cargar resultados desde disco (uploads/results/{job_id}.json)
        try:
            results_dir = RESULTS_DIR if 'RESULTS_DIR' in globals() else os.path.join(UPLOAD_DIR, 'results')
            json_path = os.path.join(results_dir, f"{job_id}.json")
            if os.path.exists(json_path):
                with open(json_path, 'r', encoding='utf-8') as rf:
                    disk_result = json.load(rf)
                # Normalizar una respuesta compatible con el esquema de status.
                # Aceptar archivos que usen 'text' o 'transcription', y 'segments' o 'diarization'.
                text_val = disk_result.get('text') or disk_result.get('transcription') or ''
                segments_val = disk_result.get('segments') or disk_result.get('diarization') or None
                response = {
                    "job_id": job_id,
                    "state": disk_result.get('state', 'done'),
                    "progress": disk_result.get('progress', 100),
                    "message": disk_result.get('message', 'completado (desde disco)'),
                    "eta_seconds": 0,
                    "result": {
                        "text": text_val,
                        "segments": segments_val,
                        "diarization": disk_result.get('diarization') or None,
                        "model": disk_result.get('whisper_model') or disk_result.get('model')
                    },
                    "error": disk_result.get('error') if isinstance(disk_result.get('error'), str) else None
                }
                return response
        except Exception:
            # Si falla la carga desde disco, continuar y devolver 404 abajo
            pass

        raise HTTPException(status_code=404, detail="Job no encontrado")

    remaining = _estimate_remaining(job)
    response = {
        "job_id": job_id,
        "state": job.get("state"),
        "progress": job.get("progress"),
        "message": job.get("message"),
        "eta_seconds": remaining,
        "result": job.get("result"),
        "error": job.get("error")
    }
    return response



def get_transcriber():
    global transcriber
    if transcriber is None:
        logger.info(f"Cargando modelo Whisper: {WHISPER_MODEL}")
        try:
            transcriber = AudioTranscriber(model_name=WHISPER_MODEL)
        except ModuleNotFoundError as e:
            # Rethrow with a clearer message for handlers
            logger.error(f"Dependencia faltante al inicializar transcriber: {e}")
            raise RuntimeError("Dependencia requerida no está instalada: torch. Instala las dependencias pesadas o usa la imagen 'full' que incluya PyTorch.")
        except Exception:
            # propagate other errors
            raise
    return transcriber


def get_diarizer():
    global diarizer
    if diarizer is None:
        logger.info("Cargando pipeline de diarización")
        diarizer = SpeakerDiarizer(hf_token=HF_TOKEN)
    return diarizer


def get_video_converter():
    global video_converter
    if video_converter is None:
        logger.info("Inicializando convertidor de video")
        video_converter = VideoConverter(output_dir=UPLOAD_DIR)
    return video_converter


def _save_docx(text: str, path: str):
    try:
        from docx import Document
    except Exception:
        raise RuntimeError("python-docx no está instalado. Instálalo con: pip install python-docx")
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(path)


def _save_pdf(text: str, path: str):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except Exception:
        raise RuntimeError("reportlab no está instalado. Instálalo con: pip install reportlab")
    width, height = letter
    c = canvas.Canvas(path, pagesize=letter)
    y = height - 72
    line_height = 12
    for line in text.splitlines():
        if y < 72:
            c.showPage()
            y = height - 72
        # Truncar líneas muy largas para evitar overflow
        c.drawString(72, y, line[:1000])
        y -= line_height
    c.save()


app.mount("/static", StaticFiles(directory=str(WEB_DIR)), name="static")


class TranscriptionResponse(BaseModel):
    job_id: str
    text: str
    language: str
    segments: List[dict]
    statistics: Optional[dict] = None
    formatted_transcript: Optional[str] = None


class HealthResponse(BaseModel):
    status: str
    timestamp: str
    whisper_model: str
    device: str


@app.get("/", response_class=HTMLResponse, tags=["General"])
async def root():
    html_file = WEB_DIR / "index.html"
    if html_file.exists():
        return FileResponse(html_file)
    else:
        return HTMLResponse(
            content="<h1>Error: No se encontró index.html</h1>",
            status_code=404
        )


@app.get("/api", tags=["General"])
async def api_info():
    return {
        "message": "API de Transcripción de Audio con Diarización",
        "version": "1.0.0",
        "endpoints": {
            "health": "/health",
            "transcribe": "/transcribe",
            "transcribe_with_diarization": "/transcribe-diarize",
            "docs": "/docs"
        }
    }


@app.get("/health", response_model=HealthResponse, tags=["General"])
async def health_check():
    # No forzar la carga del modelo aquí (evita descargar modelos grandes al consultar /health)
    device = "unknown"
    if transcriber is not None:
        try:
            device = str(transcriber.device)
        except Exception:
            device = "unknown"

    return HealthResponse(
        status="healthy",
        timestamp=datetime.now().isoformat(),
        whisper_model=WHISPER_MODEL,
        device=device
    )


@app.get("/debug", tags=["General"])
async def debug_info():
    info = {
        "status": "ok",
        "python_version": sys.version,
        "whisper_model": WHISPER_MODEL,
        "hf_token_configured": bool(HF_TOKEN),
        "hf_token_length": len(HF_TOKEN) if HF_TOKEN else 0,
        "upload_dir": UPLOAD_DIR,
        "upload_dir_exists": os.path.exists(UPLOAD_DIR),
        "web_dir": str(WEB_DIR),
        "web_dir_exists": WEB_DIR.exists(),
        "max_file_size_mb": MAX_FILE_SIZE / 1024 / 1024
    }

    try:
        import importlib
        torch = importlib.import_module('torch')
        info.update({
            "torch_version": getattr(torch, "__version__", "unknown"),
            "cuda_available": torch.cuda.is_available()
        })
    except Exception:
        info.update({
            "torch_version": None,
            "cuda_available": False
        })

    return info


class PresignResponse(BaseModel):
    upload_url: str
    object_url: str
    get_url: Optional[str] = None


@app.post('/presign', response_model=PresignResponse, tags=["Uploads"])
async def presign_upload(filename: str = Form(...)):
    """Genera una URL presigned para subir a DigitalOcean Spaces (S3 compatible).

    Requiere las variables de entorno: SPACES_KEY, SPACES_SECRET, SPACES_REGION, SPACES_BUCKET
    Si no existen, devuelve un error con instrucciones.
    """
    try:
        import boto3
    except Exception:
        raise HTTPException(status_code=501, detail="boto3 no está instalado en la imagen. Instala boto3 para usar presigned URLs.")

    SPACES_KEY = os.getenv('SPACES_KEY')
    SPACES_SECRET = os.getenv('SPACES_SECRET')
    SPACES_REGION = os.getenv('SPACES_REGION')
    SPACES_BUCKET = os.getenv('SPACES_BUCKET')

    if not (SPACES_KEY and SPACES_SECRET and SPACES_REGION and SPACES_BUCKET):
        raise HTTPException(status_code=400, detail="Faltan variables SPACES_KEY/SPACES_SECRET/SPACES_REGION/SPACES_BUCKET en el entorno")

    client = boto3.client(
        's3',
        region_name=SPACES_REGION,
        endpoint_url=f'https://{SPACES_REGION}.digitaloceanspaces.com',
        aws_access_key_id=SPACES_KEY,
        aws_secret_access_key=SPACES_SECRET
    )

    key = f"uploads/{uuid.uuid4()}_{os.path.basename(filename)}"
    try:
        url = client.generate_presigned_url(
            'put_object',
            Params={'Bucket': SPACES_BUCKET, 'Key': key, 'ACL': 'private'},
            ExpiresIn=3600
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando presigned URL: {e}")

    object_url = f"https://{SPACES_BUCKET}.{SPACES_REGION}.digitaloceanspaces.com/{key}"
    # Generar también una URL presignada para GET para que el worker pueda descargar sin hacer público el objeto
    try:
        get_url = client.generate_presigned_url(
            'get_object',
            Params={'Bucket': SPACES_BUCKET, 'Key': key},
            ExpiresIn=3600
        )
    except Exception:
        get_url = None

    return PresignResponse(upload_url=url, object_url=object_url, get_url=get_url)


class JobCreate(BaseModel):
    source_url: str
    task: str = 'transcribe-diarize'  # 'transcribe' or 'transcribe-diarize'
    whisper_model: Optional[str] = None
    num_speakers: Optional[int] = None


@app.post('/jobs/create', tags=["Jobs"])
async def create_job_from_url(payload: JobCreate):
    """Crear un job a partir de una URL accesible públicamente (ej: Spaces pre-signed object URL).

    El worker en segundo plano detectará el job y lo procesará.
    """
    job_id = str(uuid.uuid4())
    meta = payload.dict()
    meta.update({
        'job_id': job_id,
        'created_at': datetime.utcnow().isoformat()
    })

    job_file = os.path.join(JOBS_DIR, f"{job_id}.json")
    try:
        with open(job_file, 'w', encoding='utf-8') as f:
            json.dump(meta, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo encolar el job: {e}")

    _create_job(job_id, meta=meta)
    _update_job(job_id, state='queued', message='en cola', progress=0)

    return { 'job_id': job_id, 'status': 'queued' }


@app.post("/transcribe", tags=["Transcription"])
async def transcribe_audio(
    file: UploadFile = File(...),
    language: Optional[str] = Form(None),
    task: str = Form("transcribe"),
    output_format: str = Form("text"),
    download_format: Optional[str] = Form(None),
    async_process: bool = Form(False),
    background_tasks: BackgroundTasks = None
):
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"Archivo demasiado grande. Máximo: {MAX_FILE_SIZE / 1024 / 1024}MB"
        )

    job_id = str(uuid.uuid4())
    file_extension = os.path.splitext(file.filename)[1]
    temp_path = os.path.join(UPLOAD_DIR, f"{job_id}{file_extension}")

    try:
        async with aiofiles.open(temp_path, 'wb') as f:
            await f.write(content)

        logger.info(f"Procesando transcripción para job_id: {job_id}")

        # Si se solicita procesamiento asíncrono, encolar y devolver el job_id
        if async_process and background_tasks is not None:
            _create_job(job_id, meta={"type": "transcribe", "filename": file.filename})
            _update_job(job_id, state="queued", message="en cola", progress=0)
            background_tasks.add_task(run_transcription_job, job_id, temp_path, language, task, download_format)
            return {"job_id": job_id, "status": "queued"}

        # Procesamiento síncrono (comportamiento previo)
        transcriber_instance = get_transcriber()
        result = transcriber_instance.transcribe(
            temp_path,
            language=language,
            task=task
        )

        # Manejar descarga en DOCX o PDF
        if download_format in ("docx", "pdf"):
            out_name = f"{job_id}.{download_format}"
            out_path = os.path.join(UPLOAD_DIR, out_name)
            text_for_export = result.get("text", "")
            if download_format == "docx":
                _save_docx(text_for_export, out_path)
            else:
                _save_pdf(text_for_export, out_path)

            if background_tasks is not None:
                background_tasks.add_task(lambda p: os.remove(p) if os.path.exists(p) else None, out_path)

            return FileResponse(out_path, media_type="application/octet-stream", filename=out_name)

        return {
            "job_id": job_id,
            "text": result["text"],
            "language": result["language"],
            "segments": result["segments"],
            "model": result["model"]
        }

    except Exception as e:
        logger.error(f"Error en transcripción: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Error al transcribir: {type(e).__name__}: {str(e)}"
        )

    finally:
        # Si fue encolado para background, no eliminar aquí: lo hará el worker cuando termine
        if not async_process and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception as e:
                logger.warning(f"No se pudo eliminar archivo temporal: {e}")


@app.post("/transcribe-diarize", tags=["Transcription"])
async def transcribe_with_diarization(
    file: UploadFile = File(...),
    language: Optional[str] = Form(None),
    num_speakers: Optional[int] = Form(None),
    min_speakers: Optional[int] = Form(None),
    max_speakers: Optional[int] = Form(None),
    output_format: str = Form("text"),
    download_format: Optional[str] = Form(None),
    async_process: bool = Form(False),
    background_tasks: BackgroundTasks = None
):
    """
    Transcribe audio e identifica hablantes.
    
    Args:
        file: Archivo de audio
        language: Código de idioma
        num_speakers: Número exacto de hablantes (opcional)
        min_speakers: Número mínimo de hablantes (opcional)
        max_speakers: Número máximo de hablantes (opcional)
        output_format: Formato de salida ('text', 'detailed', 'srt')
    
    Returns:
        Transcripción con identificación de hablantes
    """
    if not HF_TOKEN:
        logger.error("HF_TOKEN no configurado")
        raise HTTPException(
            status_code=503,
            detail="HF_TOKEN no configurado. Se requiere para diarización."
        )
    
    # Leer contenido
    try:
        content = await file.read()
    except Exception as e:
        logger.error(f"Error al leer archivo: {e}")
        raise HTTPException(
            status_code=400,
            detail=f"Error al leer el archivo: {str(e)}"
        )
    
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"Archivo demasiado grande. Máximo: {MAX_FILE_SIZE / 1024 / 1024}MB"
        )
    
    if len(content) == 0:
        raise HTTPException(
            status_code=400,
            detail="El archivo está vacío"
        )
    
    job_id = str(uuid.uuid4())
    file_extension = os.path.splitext(file.filename)[1] if file.filename else ".webm"
    
    # Si no tiene extensión o es webm, convertir a wav
    if not file_extension or file_extension.lower() in ['.webm', '.ogg']:
        file_extension = '.wav'
    
    temp_path = os.path.join(UPLOAD_DIR, f"{job_id}{file_extension}")
    
    logger.info(f"=== INICIANDO PROCESAMIENTO ===")
    logger.info(f"Job ID: {job_id}")
    logger.info(f"Archivo original: {file.filename}")
    logger.info(f"Tipo MIME: {file.content_type}")
    logger.info(f"Tamaño: {len(content) / 1024:.2f} KB")
    logger.info(f"Ruta temporal: {temp_path}")
    
    try:
        async with aiofiles.open(temp_path, 'wb') as f:
            await f.write(content)
        
        logger.info(f"Procesando transcripción + diarización para job_id: {job_id}")
        logger.info(f"Archivo guardado: {temp_path}")
        logger.info(f"Tamaño: {len(content) / 1024:.2f} KB")
        
        # Si se solicita procesamiento asíncrono, encolar y devolver job_id
        if async_process:
            # El archivo ya está guardado en temp_path, no necesitamos copiarlo
            # Solo aseguramos que la ruta es la correcta para el worker
            audio_file_path = temp_path
            logger.info(f"[JOB {job_id}] Archivo listo para procesamiento: {audio_file_path}")
            
            _create_job(job_id, meta={"type": "transcribe-diarize", "filename": file.filename})
            _update_job(job_id, state="queued", message="en cola", progress=0)
            
            def _run_diarize_job():
                try:
                    _update_job(job_id, state="running", message="iniciando", progress=5)
                    
                    # 1. Transcribir
                    logger.info(f"[JOB {job_id}] Cargando modelo Whisper...")
                    transcriber_instance = get_transcriber()
                    
                    logger.info(f"[JOB {job_id}] Transcribiendo audio...")
                    trans_result = transcriber_instance.transcribe_with_timestamps(audio_file_path, language=language)
                    _update_job(job_id, progress=40, message="transcripción completada")
                    
                    # 2. Diarizar
                    logger.info(f"[JOB {job_id}] Cargando diarizador...")
                    diarizer_instance = get_diarizer()
                    
                    logger.info(f"[JOB {job_id}] Diarizando...")
                    dia_segments = diarizer_instance.diarize(audio_file_path, num_speakers=num_speakers, min_speakers=min_speakers, max_speakers=max_speakers)
                    _update_job(job_id, progress=70, message="diarización completada")
                    
                    # 3. Combinar
                    aligned_segments = align_transcription_with_diarization(trans_result, dia_segments)
                    formatted_text = format_transcript(aligned_segments, output_format)
                    statistics = get_speaker_statistics(aligned_segments)
                    
                    # Persistir resultados
                    results_dir = RESULTS_DIR if 'RESULTS_DIR' in globals() else os.path.join(UPLOAD_DIR, 'results')
                    os.makedirs(results_dir, exist_ok=True)
                    job_result = {
                        "job_id": job_id,
                        "text": formatted_text,
                        "segments": aligned_segments,
                        "statistics": statistics,
                        "num_speakers": len(statistics),
                        "output_format": output_format,
                        "timestamp": datetime.utcnow().isoformat(),
                        "whisper_model": WHISPER_MODEL
                    }
                    json_path = os.path.join(results_dir, f"{job_id}.json")
                    with open(json_path, 'w', encoding='utf-8') as rf:
                        json.dump(job_result, rf, ensure_ascii=False, indent=2)
                    
                    txt_path = os.path.join(results_dir, f"{job_id}.txt")
                    with open(txt_path, 'w', encoding='utf-8') as tf:
                        tf.write(formatted_text)
                    
                    # Actualizar job en memoria
                    _update_job(job_id, state='done', message='completado', progress=100, result={"text": formatted_text, "segments": aligned_segments, "statistics": statistics, "num_speakers": len(statistics)})
                    
                except Exception as e:
                    logger.exception(f"[JOB {job_id}] Error: {e}")
                    _update_job(job_id, state='error', error=str(e), message='error')
                finally:
                    # Limpiar archivo de audio
                    if os.path.exists(audio_file_path):
                        try:
                            os.remove(audio_file_path)
                            logger.info(f"[JOB {job_id}] Archivo eliminado: {audio_file_path}")
                        except Exception as e_rm:
                            logger.warning(f"[JOB {job_id}] No se pudo eliminar archivo: {e_rm}")
            
            # Usar BackgroundTasks si está disponible, sino threading
            if background_tasks is not None:
                background_tasks.add_task(_run_diarize_job)
            else:
                import threading
                thread = threading.Thread(target=_run_diarize_job, daemon=True)
                thread.start()
            
            return {"job_id": job_id, "status": "queued"}
        
        # Procesamiento síncrono (comportamiento anterior)
        # 1. Transcribir
        logger.info("Paso 1/5: Cargando modelo Whisper...")
        transcriber_instance = get_transcriber()
        
        logger.info("Paso 2/5: Transcribiendo audio...")
        trans_result = transcriber_instance.transcribe_with_timestamps(
            temp_path,
            language=language
        )
        logger.info(f"Transcripción completada: {len(trans_result)} segmentos")
        
        # 2. Diarizar
        logger.info("Paso 3/5: Cargando modelo de diarización...")
        try:
            diarizer_instance = get_diarizer()
        except Exception as e:
            logger.error(f"Error al cargar diarizador: {e}", exc_info=True)
            raise HTTPException(
                status_code=500,
                detail=f"Error al cargar modelo de diarización: {str(e)}. Verifica tu HF_TOKEN y conexión a internet."
            )
        
        logger.info("Paso 4/5: Identificando hablantes...")
        dia_segments = diarizer_instance.diarize(
            temp_path,
            num_speakers=num_speakers,
            min_speakers=min_speakers,
            max_speakers=max_speakers
        )
        logger.info(f"Diarización completada: {len(dia_segments)} segmentos")
        
        # 3. Combinar
        logger.info("Paso 5/5: Combinando resultados...")
        aligned_segments = align_transcription_with_diarization(
            trans_result,
            dia_segments
        )
        
        # 4. Formatear
        formatted_text = format_transcript(aligned_segments, output_format)
        
        # 5. Estadísticas
        statistics = get_speaker_statistics(aligned_segments)

        # Persistir resultados en disco para que /status pueda localizarlos posteriormente
        try:
            results_dir = RESULTS_DIR if 'RESULTS_DIR' in globals() else os.path.join(UPLOAD_DIR, 'results')
            # ensure_upload_dirs ya deberda haber creado los directorios en startup
            os.makedirs(results_dir, exist_ok=True)

            job_result = {
                "job_id": job_id,
                "text": formatted_text,
                "segments": aligned_segments,
                "statistics": statistics,
                "num_speakers": len(statistics),
                "output_format": output_format,
                "timestamp": datetime.utcnow().isoformat(),
                "whisper_model": WHISPER_MODEL
            }

            json_path = os.path.join(results_dir, f"{job_id}.json")
            with open(json_path, 'w', encoding='utf-8') as rf:
                json.dump(job_result, rf, ensure_ascii=False, indent=2)

            # Guardar también la versión de texto simple para descarga rápida
            txt_path = os.path.join(results_dir, f"{job_id}.txt")
            try:
                with open(txt_path, 'w', encoding='utf-8') as tf:
                    tf.write(formatted_text)
            except Exception:
                pass

            logger.info(f"Resultados guardados en: {json_path}")

            # Actualizar store en memoria si existe (útil justo después para /status en la misma instancia)
            try:
                _create_job(job_id, meta={"type": "transcribe-diarize", "filename": file.filename})
                _update_job(job_id, state='done', message='completado', progress=100, result={"text": formatted_text, "segments": aligned_segments, "statistics": statistics})
            except Exception:
                # No crítico si falla actualizar memoria
                pass
        except Exception as e:
            logger.warning(f"No se pudieron persistir los resultados a disco: {e}")

        # Manejar descarga en DOCX o PDF
        if download_format in ("docx", "pdf"):
            out_name = f"{job_id}.{download_format}"
            out_path = os.path.join(UPLOAD_DIR, out_name)
            if download_format == "docx":
                _save_docx(formatted_text, out_path)
            else:
                _save_pdf(formatted_text, out_path)

            if background_tasks is not None:
                background_tasks.add_task(lambda p: os.remove(p) if os.path.exists(p) else None, out_path)

            return FileResponse(out_path, media_type="application/octet-stream", filename=out_name)

        return {
            "job_id": job_id,
            "text": formatted_text,
            "segments": aligned_segments,
            "statistics": statistics,
            "num_speakers": len(statistics),
            "output_format": output_format
        }
        
    except Exception as e:
        logger.error(f"Error en transcripción + diarización: {e}", exc_info=True)
        raise HTTPException(
            status_code=500, 
            detail=f"Error al procesar: {type(e).__name__}: {str(e)}"
        )
    
    finally:
        
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception as e:
                logger.warning(f"No se pudo eliminar archivo temporal: {e}")


@app.post("/convert-video", tags=["Video"])
async def convert_video_to_audio(
    file: UploadFile = File(...),
    bitrate: str = Form("192k"),
    sample_rate: int = Form(44100)
):
    """
    Convierte un archivo de video a MP3 (solo audio).
    
    Args:
        file: Archivo de video (mp4, avi, mov, mkv, etc.)
        bitrate: Bitrate del audio (ej: "128k", "192k", "320k")
        sample_rate: Frecuencia de muestreo (44100, 48000)
    
    Returns:
        Archivo MP3 con el audio extraído
    """
    if not is_video_file(file.filename):
        raise HTTPException(
            status_code=400,
            detail="El archivo debe ser un video (mp4, avi, mov, mkv, etc.)"
        )
    
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"Archivo demasiado grande. Máximo: {MAX_FILE_SIZE / 1024 / 1024}MB"
        )
    
    job_id = str(uuid.uuid4())
    file_extension = os.path.splitext(file.filename)[1]
    temp_video_path = os.path.join(UPLOAD_DIR, f"{job_id}{file_extension}")
    
    try:
        async with aiofiles.open(temp_video_path, 'wb') as f:
            await f.write(content)
        
        logger.info(f"Convirtiendo video a MP3 para job_id: {job_id}")
        
        converter = get_video_converter()
        mp3_filename = f"{job_id}.mp3"
        mp3_path = converter.convert_video_to_mp3(
            temp_video_path,
            output_filename=mp3_filename,
            bitrate=bitrate,
            sample_rate=sample_rate
        )
        
        # Obtener información del video
        video_info = converter.get_video_info(temp_video_path)
        
        return {
            "job_id": job_id,
            "message": "Video convertido exitosamente a MP3",
            "mp3_filename": mp3_filename,
            "mp3_path": mp3_path,
            "video_info": {
                "duration": video_info.get("duration", 0),
                "had_audio": video_info.get("has_audio", False),
                "original_format": video_info.get("format", "unknown")
            },
            "audio_settings": {
                "bitrate": bitrate,
                "sample_rate": sample_rate
            }
        }
        
    except Exception as e:
        logger.error(f"Error al convertir video: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))
    
    finally:
        if os.path.exists(temp_video_path):
            try:
                os.remove(temp_video_path)
                logger.info(f"Archivo temporal de video eliminado: {temp_video_path}")
            except Exception as e:
                logger.warning(f"No se pudo eliminar video temporal: {e}")


@app.post("/convert-and-transcribe", tags=["Video"])
async def convert_video_and_transcribe(
    file: UploadFile = File(...),
    language: Optional[str] = Form(None),
    num_speakers: Optional[int] = Form(None),
    min_speakers: Optional[int] = Form(None),
    max_speakers: Optional[int] = Form(None),
    output_format: str = Form("text"),
    bitrate: str = Form("192k"),
    background_tasks: BackgroundTasks = None
):
    """
    Convierte un video a audio y lo transcribe con diarización de forma ASÍNCRONA.
    Retorna inmediatamente un job_id para hacer polling del progreso.
    
    Args:
        file: Archivo de video
        language: Código de idioma
        num_speakers: Número exacto de hablantes (opcional)
        min_speakers: Número mínimo de hablantes (opcional)
        max_speakers: Número máximo de hablantes (opcional)
        output_format: Formato de salida ('text', 'detailed', 'srt')
        bitrate: Bitrate del audio
    
    Returns:
        job_id para hacer polling en /status/{job_id}
    """
    if not HF_TOKEN:
        raise HTTPException(
            status_code=503,
            detail="HF_TOKEN no configurado. Se requiere para diarización."
        )
    
    if not is_video_file(file.filename):
        raise HTTPException(
            status_code=400,
            detail="El archivo debe ser un video"
        )
    
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"Archivo demasiado grande. Máximo: {MAX_FILE_SIZE / 1024 / 1024}MB"
        )
    
    job_id = str(uuid.uuid4())
    file_extension = os.path.splitext(file.filename)[1]
    temp_video_path = os.path.join(UPLOAD_DIR, f"{job_id}{file_extension}")
    
    try:
        async with aiofiles.open(temp_video_path, 'wb') as f:
            await f.write(content)
        
        logger.info(f"Video guardado para procesamiento async: {job_id}")
        
        # Crear job y procesarlo en segundo plano
        _create_job(job_id, meta={"type": "convert-and-transcribe", "filename": file.filename})
        _update_job(job_id, state="queued", message="Video en cola para conversión", progress=0)
        
        def _process_video_job():
            mp3_path = None
            try:
                _update_job(job_id, state="running", message="Convirtiendo video a audio...", progress=5)
                
                # 1. Convertir a MP3
                converter = get_video_converter()
                mp3_filename = f"{job_id}.mp3"
                mp3_path = converter.convert_video_to_mp3(
                    temp_video_path,
                    output_filename=mp3_filename,
                    bitrate=bitrate
                )
                _update_job(job_id, progress=15, message="Conversión completada, cargando modelo Whisper...")
                
                # 2. Transcribir
                transcriber_instance = get_transcriber()
                _update_job(job_id, progress=20, message="Transcribiendo audio...")
                trans_result = transcriber_instance.transcribe_with_timestamps(
                    mp3_path,
                    language=language
                )
                _update_job(job_id, progress=55, message="Transcripción completada")
                
                # 3. Diarizar
                _update_job(job_id, progress=60, message="Cargando modelo de diarización...")
                diarizer_instance = get_diarizer()
                _update_job(job_id, progress=65, message="Identificando hablantes...")
                dia_segments = diarizer_instance.diarize(
                    mp3_path,
                    num_speakers=num_speakers,
                    min_speakers=min_speakers,
                    max_speakers=max_speakers
                )
                _update_job(job_id, progress=90, message="Diarización completada")
                
                # 4. Combinar y formatear
                _update_job(job_id, progress=93, message="Combinando resultados...")
                aligned_segments = align_transcription_with_diarization(trans_result, dia_segments)
                aligned_segments = renumber_speakers(aligned_segments)
                formatted_text = format_transcript(aligned_segments, output_format)
                statistics = get_speaker_statistics(aligned_segments)
                
                _update_job(job_id, progress=98, message="Guardando resultados...")
                
                # Guardar resultados
                result = {
                    "text": formatted_text,
                    "segments": aligned_segments,
                    "statistics": statistics,
                    "num_speakers": len(statistics),
                    "output_format": output_format,
                    "source": "video_conversion"
                }
                
                results_dir = os.path.join(UPLOAD_DIR, 'results')
                os.makedirs(results_dir, exist_ok=True)
                json_path = os.path.join(results_dir, f"{job_id}.json")
                with open(json_path, 'w', encoding='utf-8') as jf:
                    json.dump(result, jf, ensure_ascii=False, indent=2)
                
                _update_job(job_id, progress=100, state="done", message="Completado", result=result)
                logger.info(f"[JOB {job_id}] Video procesado exitosamente")
                
            except Exception as e:
                logger.exception(f"[JOB {job_id}] Error procesando video: {e}")
                _update_job(job_id, state="error", error=str(e), message=f"Error: {str(e)}")
            finally:
                # Limpiar archivos temporales
                for temp_file in [temp_video_path, mp3_path]:
                    if temp_file and os.path.exists(temp_file):
                        try:
                            os.remove(temp_file)
                            logger.info(f"[JOB {job_id}] Eliminado: {temp_file}")
                        except Exception as e:
                            logger.warning(f"[JOB {job_id}] No se pudo eliminar {temp_file}: {e}")
        
        # Ejecutar en segundo plano
        if background_tasks:
            background_tasks.add_task(_process_video_job)
        else:
            import threading
            t = threading.Thread(target=_process_video_job, daemon=True)
            t.start()
        
        return {
            "job_id": job_id,
            "message": "Video encolado para procesamiento",
            "status_url": f"/status/{job_id}"
        }
        
    except Exception as e:
        logger.error(f"Error al encolar video: {e}", exc_info=True)
        # Limpiar archivo temporal si falla
        if os.path.exists(temp_video_path):
            try:
                os.remove(temp_video_path)
            except Exception:
                pass
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/models", tags=["General"])
async def get_models_info():
    """Información sobre los modelos disponibles"""
    return {
        "whisper_models": ["tiny", "base", "small", "medium", "large", "large-v2", "large-v3"],
        "current_whisper_model": WHISPER_MODEL,
        "diarization_model": "pyannote/speaker-diarization-3.1",
        "supported_languages": [
            "es", "en", "fr", "de", "it", "pt", "ru", "zh", "ja", "ko"
        ],
        "supported_video_formats": [
            "mp4", "avi", "mov", "mkv", "flv", "wmv", "webm", "m4v", "mpg", "mpeg", "3gp", "ogv"
        ]
    }


@app.post('/debug/create_test_job', tags=["General"])
async def create_test_job(duration_seconds: int = 10, background_tasks: BackgroundTasks = None):
    """Crear un job de prueba que avanza automáticamente para testear el polling UI.

    Args:
        duration_seconds: Tiempo total que tomará el job (aprox.)
    Returns:
        job_id
    """
    job_id = str(uuid.uuid4())
    _create_job(job_id, meta={"type": "test"})
    _update_job(job_id, state="queued", message="en cola", progress=0)

    def _runner(jid: str, total: int):
        try:
            _update_job(jid, state="running", message="iniciando", progress=1)
            steps = max(4, int(total))
            for i in range(1, steps + 1):
                time.sleep(total / steps)
                prog = int((i / steps) * 100)
                _update_job(jid, progress=prog, message=f"progreso {prog}%")
            # Resultado de ejemplo
            res = {
                "text": "Este es un resultado de prueba.",
                "language": "es",
                "segments": [],
                "model": "test-sim"
            }
            _update_job(jid, progress=100, state="done", message="completado", result=res)
        except Exception as e:
            _update_job(jid, state="error", error=str(e), message="error")

    # Ejecutar en background
    if background_tasks is not None:
        background_tasks.add_task(_runner, job_id, duration_seconds)
    else:
        # Si no hay background tasks (ej: llamado desde TestClient sin), lanzar hilo
        import threading
        t = threading.Thread(target=_runner, args=(job_id, duration_seconds), daemon=True)
        t.start()

    return {"job_id": job_id}



@app.post('/debug/remove_job_file', tags=["General"])
async def debug_remove_job_file(job_id: str):
    """Eliminar o mover el archivo de job en uploads/jobs/ por job_id.

    Útil para demorar o limpiar jobs problemáticos sin acceso directo al FS del contenedor.
    """
    job_file = os.path.join(JOBS_DIR, f"{job_id}.json")
    failed_dir = os.path.join(JOBS_DIR, 'failed')
    if not os.path.exists(job_file):
        return {"ok": False, "reason": "job_file_not_found", "path": job_file}

    try:
        os.makedirs(failed_dir, exist_ok=True)
        # mover a failed con sufijo timestamp
        ts = datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')
        dest = os.path.join(failed_dir, f"{job_id}.{ts}.failed")
        os.replace(job_file, dest)
        return {"ok": True, "moved_to": dest}
    except Exception as e:
        logger.exception(f"Error moviendo job file {job_file}: {e}")
        return {"ok": False, "reason": str(e)}


# Endpoints de descarga en formatos PDF y Word
@app.post('/download/pdf', tags=["Downloads"])
async def download_pdf(result: dict):
    """
    Genera un PDF a partir del resultado de transcripción.
    
    Args:
        result: Objeto JSON con el resultado de la transcripción
    
    Returns:
        Archivo PDF descargable
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas as pdf_canvas
        from reportlab.lib.units import inch
        import io
        
        # Extraer texto del resultado
        def extract_text(r):
            if isinstance(r, str):
                return r
            if isinstance(r, dict):
                if 'text' in r:
                    return r['text']
                if 'transcription' in r:
                    return r['transcription']
                if 'segments' in r and isinstance(r['segments'], list):
                    return '\n\n'.join(seg.get('text', '') for seg in r['segments'])
            return str(r)
        
        text = extract_text(result)
        
        # Crear PDF en memoria
        buffer = io.BytesIO()
        c = pdf_canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        # Título
        c.setFont("Helvetica-Bold", 16)
        c.drawString(inch, height - inch, "Transcripción de Audio")
        
        # Fecha
        c.setFont("Helvetica", 10)
        c.drawString(inch, height - inch - 0.3*inch, f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        
        # Contenido
        c.setFont("Helvetica", 11)
        y_position = height - inch - 0.7*inch
        
        # Dividir texto en líneas y páginas
        lines = text.split('\n')
        for line in lines:
            # Wrap texto largo
            if len(line) > 80:
                words = line.split()
                current_line = ""
                for word in words:
                    test_line = current_line + word + " "
                    if len(test_line) > 80:
                        if current_line:
                            c.drawString(inch, y_position, current_line.strip())
                            y_position -= 0.2*inch
                            if y_position < inch:
                                c.showPage()
                                c.setFont("Helvetica", 11)
                                y_position = height - inch
                        current_line = word + " "
                    else:
                        current_line = test_line
                
                if current_line:
                    c.drawString(inch, y_position, current_line.strip())
                    y_position -= 0.2*inch
            else:
                c.drawString(inch, y_position, line)
                y_position -= 0.2*inch
            
            # Nueva página si es necesario
            if y_position < inch:
                c.showPage()
                c.setFont("Helvetica", 11)
                y_position = height - inch
        
        # Estadísticas si existen
        if isinstance(result, dict) and 'statistics' in result:
            c.showPage()
            c.setFont("Helvetica-Bold", 14)
            y_position = height - inch
            c.drawString(inch, y_position, "Estadísticas de Hablantes")
            y_position -= 0.4*inch
            
            c.setFont("Helvetica", 10)
            for speaker, stats in result['statistics'].items():
                c.drawString(inch, y_position, f"{speaker}:")
                y_position -= 0.2*inch
                c.drawString(inch + 0.3*inch, y_position, f"Tiempo total: {stats.get('total_time', 0):.1f}s")
                y_position -= 0.15*inch
                c.drawString(inch + 0.3*inch, y_position, f"Palabras: {stats.get('total_words', 0)}")
                y_position -= 0.15*inch
                c.drawString(inch + 0.3*inch, y_position, f"Participación: {stats.get('time_percentage', 0):.1f}%")
                y_position -= 0.3*inch
        
        c.save()
        buffer.seek(0)
        
        from fastapi.responses import Response
        return Response(
            content=buffer.getvalue(),
            media_type='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="transcripcion_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
            }
        )
    
    except Exception as e:
        logger.error(f"Error generando PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando PDF: {str(e)}")


@app.post('/download/word', tags=["Downloads"])
async def download_word(result: dict):
    """
    Genera un documento Word (.docx) a partir del resultado de transcripción.
    
    Args:
        result: Objeto JSON con el resultado de la transcripción
    
    Returns:
        Archivo DOCX descargable
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        # Extraer texto del resultado
        def extract_text(r):
            if isinstance(r, str):
                return r
            if isinstance(r, dict):
                if 'text' in r:
                    return r['text']
                if 'transcription' in r:
                    return r['transcription']
                if 'segments' in r and isinstance(r['segments'], list):
                    return '\n\n'.join(seg.get('text', '') for seg in r['segments'])
            return str(r)
        
        text = extract_text(result)
        
        # Crear documento
        doc = Document()
        
        # Título
        title = doc.add_heading('Transcripción de Audio', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Fecha
        date_para = doc.add_paragraph()
        date_para.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}").italic = True
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Separador
        doc.add_paragraph()
        
        # Contenido de la transcripción
        doc.add_heading('Transcripción', 1)
        
        # Agregar texto por párrafos
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.style.font.size = Pt(11)
        
        # Estadísticas si existen
        if isinstance(result, dict) and 'statistics' in result:
            doc.add_page_break()
            doc.add_heading('Estadísticas de Hablantes', 1)
            
            for speaker, stats in result['statistics'].items():
                doc.add_heading(speaker, 2)
                doc.add_paragraph(f"Tiempo total: {stats.get('total_time', 0):.1f} segundos")
                doc.add_paragraph(f"Palabras: {stats.get('total_words', 0)}")
                doc.add_paragraph(f"Participación: {stats.get('time_percentage', 0):.1f}%")
                doc.add_paragraph(f"Segmentos: {stats.get('segment_count', 0)}")
                doc.add_paragraph()
        
        # Guardar en memoria
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        from fastapi.responses import Response
        return Response(
            content=buffer.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="transcripcion_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx"'
            }
        )
    
    except Exception as e:
        logger.error(f"Error generando documento Word: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando documento Word: {str(e)}")


# Manejo de errores
@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    logger.error(f"Error no manejado: {exc}")
    return JSONResponse(
        status_code=500,
        content={"detail": "Error interno del servidor", "error": str(exc)}
    )


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8888))
    host = os.getenv("HOST", "0.0.0.0")
    
    print("=" * 60)
    print("🎤 API de Transcripción de Audio con Diarización")
    print("=" * 60)
    print(f"📡 Servidor: http://{host}:{port}")
    print(f"📚 Documentación: http://{host}:{port}/docs")
    print(f"🌐 Interfaz Web: http://{host}:{port}")
    print(f"🤖 Modelo Whisper: {WHISPER_MODEL}")
    print(f"🔑 HF Token: {'✅ Configurado' if HF_TOKEN else '❌ No configurado'}")
    print("=" * 60)
    print("Presiona Ctrl+C para detener el servidor\n")
    
    uvicorn.run(
        app,  # Pasar la app directamente (más rápido, sin reload)
        host=host, 
        port=port,
        log_level="info"
    )
